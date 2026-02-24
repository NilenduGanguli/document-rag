#!/usr/bin/env python3
"""
prepare_test_documents.py
=========================
Downloads real public examples of every document type described in
US_Public_Company_Document_Guide.docx and optionally runs ingestion smoke
tests against the document-rag API.

The catalog of documents to download is driven by:
  test-scripts/document_catalog.xlsx

Edit that file to add, remove, or change source URLs / CIKs without touching
this script.  Rows with  manual_download = TRUE  are skipped during auto-
download; they are listed in  test-data/needed_files.md  so you can place them
manually.

Sources (determined per-row from the catalog)
----------------------------------------------
  SEC EDGAR  — rows where cik is set; primary HTM document → converted to PDF
  FDIC BankFind — rows where source_url contains "banks.data.fdic.gov"
  python-docx generator — rows where source_url == "generated"
  Image fixture — rows where source_url == "test fixture"

Usage
-----
  # Download only (no API needed)
  python test-scripts/prepare_test_documents.py --skip-ingest

  # Full smoke test (API must be running)
  docker compose up -d
  python test-scripts/prepare_test_documents.py \\
      --api-url http://localhost:8000 \\
      --customer-id test-co-001

  # Limit to specific forms
  python test-scripts/prepare_test_documents.py --forms 10-K,DEF14A --skip-ingest
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import requests

# ── Constants ──────────────────────────────────────────────────────────────────

EDGAR_SUBMISSIONS = "https://data.sec.gov/submissions"
EDGAR_ARCHIVE     = "https://www.sec.gov/Archives/edgar/data"
USER_AGENT        = "KYC-RAG-Test/1.0 test@example.com"
MAX_FILE_BYTES    = 20 * 1024 * 1024   # 20 MB hard cap
EDGAR_RATE_LIMIT  = 0.12               # ≈8 req/s  (SEC limit: 10 req/s)

SCRIPT_DIR    = Path(__file__).parent.resolve()
REPO_ROOT     = SCRIPT_DIR.parent
TEST_DATA_ROOT  = REPO_ROOT / "test-data"
CATALOG_PATH    = SCRIPT_DIR / "document_catalog.xlsx"

SEC_FILINGS_DIR = TEST_DATA_ROOT / "sec_filings"
BANKING_DIR     = TEST_DATA_ROOT / "banking"
IMAGES_DIR      = TEST_DATA_ROOT / "images"
DOCUMENTS_DIR   = TEST_DATA_ROOT / "documents"
MANIFEST_PATH   = TEST_DATA_ROOT / "manifest.json"
NEEDED_PATH     = TEST_DATA_ROOT / "needed_files.md"

# ── Catalog loader ─────────────────────────────────────────────────────────────

def load_catalog() -> List[Dict[str, Any]]:
    """
    Read document_catalog.xlsx and return a list of row dicts.
    Required columns: form_type, cik, company, description, filename,
                      source_url, manual_download, notes
    """
    try:
        import openpyxl
    except ImportError:
        print("[error] openpyxl not installed — cannot read catalog XLSX.")
        sys.exit(1)

    if not CATALOG_PATH.exists():
        print(f"[error] Catalog not found: {CATALOG_PATH}")
        sys.exit(1)

    wb = openpyxl.load_workbook(CATALOG_PATH, read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    if not rows:
        print("[error] Catalog is empty.")
        sys.exit(1)

    headers = [str(h).strip().lower() if h else "" for h in rows[0]]
    catalog = []
    for raw in rows[1:]:
        row = {headers[i]: (str(v).strip() if v is not None else "") for i, v in enumerate(raw)}
        if row.get("form_type"):   # skip blank rows
            # Normalise manual_download to bool
            row["manual_download"] = row.get("manual_download", "FALSE").upper() == "TRUE"
            catalog.append(row)
    return catalog


# ── Shared HTTP session ────────────────────────────────────────────────────────

_session = requests.Session()
_session.headers.update({"User-Agent": USER_AGENT, "Accept-Encoding": "identity"})
_last_edgar_ts = 0.0


def _edgar_get(url: str, stream: bool = False, timeout: int = 30) -> requests.Response:
    global _last_edgar_ts
    wait = EDGAR_RATE_LIMIT - (time.time() - _last_edgar_ts)
    if wait > 0:
        time.sleep(wait)
    resp = _session.get(url, stream=stream, timeout=timeout)
    _last_edgar_ts = time.time()
    resp.raise_for_status()
    return resp


# ── EDGAR helpers ──────────────────────────────────────────────────────────────

def _cik10(cik: str) -> str:
    return cik.zfill(10)


def _get_latest_filing(cik: str, form_type: str) -> Optional[Tuple[str, str]]:
    """Return (acc_nodash, primary_document) for the most recent filing, or None."""
    url = f"{EDGAR_SUBMISSIONS}/CIK{_cik10(cik)}.json"
    try:
        data = _edgar_get(url, timeout=20).json()
    except Exception as exc:
        print(f"    [warn] submissions fetch failed for CIK {cik}: {exc}")
        return None

    recent = data.get("filings", {}).get("recent", {})
    target = form_type.strip().upper()
    for form, acc, pdoc in zip(
        recent.get("form", []),
        recent.get("accessionNumber", []),
        recent.get("primaryDocument", []),
    ):
        if form.strip().upper() == target and pdoc:
            return acc.replace("-", ""), pdoc

    for extra in data.get("filings", {}).get("files", []):
        try:
            edata = _edgar_get(f"{EDGAR_SUBMISSIONS}/{extra['name']}", timeout=20).json()
            for form, acc, pdoc in zip(
                edata.get("form", []),
                edata.get("accessionNumber", []),
                edata.get("primaryDocument", []),
            ):
                if form.strip().upper() == target and pdoc:
                    return acc.replace("-", ""), pdoc
        except Exception:
            pass
    return None


def _efts_find(form_type: str) -> Optional[Tuple[str, str, str]]:
    """Return (cik, acc_nodash, company_name) from EFTS search, or None."""
    url = (
        "https://efts.sec.gov/LATEST/search-index"
        f"?q=&forms={requests.utils.quote(form_type)}"
        "&dateRange=custom&startdt=2023-01-01"
    )
    try:
        hits = _edgar_get(url, timeout=15).json().get("hits", {}).get("hits", [])
        if hits:
            src = hits[0]["_source"]
            ciks = src.get("ciks", [])
            adsh = src.get("adsh", "").replace("-", "")
            names = src.get("display_names", [""])
            if ciks and adsh:
                raw_cik = ciks[0].lstrip("0") or "0"
                company = names[0].split("(CIK")[0].strip() if names else "Unknown"
                return raw_cik, adsh, company
    except Exception as exc:
        print(f"    [warn] EFTS search failed for {form_type}: {exc}")
    return None


def _get_primary_doc_from_submissions(cik: str, acc_nodash: str) -> Optional[str]:
    try:
        data = _edgar_get(f"{EDGAR_SUBMISSIONS}/CIK{_cik10(cik)}.json", timeout=20).json()
        recent = data.get("filings", {}).get("recent", {})
        for acc, pdoc in zip(recent.get("accessionNumber", []), recent.get("primaryDocument", [])):
            if acc.replace("-", "") == acc_nodash and pdoc:
                return pdoc
        for extra in data.get("filings", {}).get("files", []):
            try:
                edata = _edgar_get(f"{EDGAR_SUBMISSIONS}/{extra['name']}", timeout=20).json()
                for acc, pdoc in zip(edata.get("accessionNumber", []), edata.get("primaryDocument", [])):
                    if acc.replace("-", "") == acc_nodash and pdoc:
                        return pdoc
            except Exception:
                pass
    except Exception:
        pass
    return None


# ── HTML → PDF converter ───────────────────────────────────────────────────────

def _to_pdf(src_path: Path) -> Optional[Path]:
    """
    Convert an HTM/HTML file to a text-based PDF using fpdf2 + html2text.
    Only HTML files are supported; returns None for other types.
    """
    if src_path.suffix.lower() not in (".htm", ".html"):
        return None

    try:
        from fpdf import FPDF
        import html2text as _h2t
    except ImportError:
        print("    [warn] fpdf2/html2text not installed — cannot convert to PDF.")
        return None

    try:
        raw = src_path.read_bytes().decode("utf-8", errors="replace")
        h = _h2t.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.body_width = 100
        text = h.handle(raw)

        if len(text) > 300_000:
            text = text[:300_000] + "\n\n[... content truncated for file size ...]"

        def _safe(s: str) -> str:
            return s.encode("latin-1", errors="replace").decode("latin-1")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=12)
        pdf.add_page()
        pdf.set_font("Helvetica", size=8)
        W = pdf.epw
        for line in text.splitlines():
            clean = _safe(line.strip())
            if not clean:
                pdf.ln(2)
                continue
            try:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(W, 4, text=clean)
            except Exception:
                pass

        pdf_path = src_path.with_suffix(".pdf")
        pdf.output(str(pdf_path))
        return pdf_path

    except Exception as exc:
        print(f"    [warn] PDF conversion failed for {src_path.name}: {exc}")
        return None


# ── EDGAR downloader ───────────────────────────────────────────────────────────

# Extensions that will be auto-converted to PDF
_HTM_EXTS = {".htm", ".html"}
# Extensions that require manual download (XML, TXT, etc.)
_MANUAL_EXTS = {".xml", ".txt", ".xsd", ".json"}


def download_from_edgar(
    form_type: str,
    cik: str,
    company: str,
    filename: str,
    save_dir: Path,
) -> Tuple[Optional[Dict[str, Any]], Optional[Dict[str, Any]]]:
    """
    Download the most recent HTM filing for `form_type` / `cik`, convert to PDF.

    Returns:
      (manifest_entry, None)          — on successful download
      (None, needed_entry)            — if primary document is XML/TXT (manual needed)
      (None, None)                    — on skip/error
    """
    # ── Already have a PDF? ───────────────────────────────────────────────────
    pdf_path = save_dir / f"{filename}.pdf"
    if pdf_path.exists():
        size = pdf_path.stat().st_size
        print(f"  [exists] {pdf_path.name} — skipping {form_type}")
        return {
            "form_type": form_type, "company": company,
            "source_url": "cached",
            "file_path": str(pdf_path.relative_to(REPO_ROOT)),
            "format": "pdf", "size_bytes": size,
            "downloaded_at": datetime.fromtimestamp(
                pdf_path.stat().st_mtime, tz=timezone.utc
            ).isoformat(),
        }, None

    # Remove any stale non-PDF remnants
    for stale in save_dir.glob(f"{filename}.*"):
        if stale.suffix.lower() != ".pdf":
            stale.unlink(missing_ok=True)

    print(f"  Fetching {form_type} ({company}) …")

    # ── Resolve accession + primary document ─────────────────────────────────
    acc_nodash: Optional[str] = None
    primary_doc: Optional[str] = None

    if cik == "0":
        result = _efts_find(form_type)
        if not result:
            print(f"    [skip] No recent {form_type} found via EFTS.")
            return None, None
        cik, acc_nodash, company = result
        primary_doc = _get_primary_doc_from_submissions(cik, acc_nodash)
        if not primary_doc:
            print(f"    [skip] Could not resolve primary document for EFTS hit.")
            return None, None
    else:
        filing = _get_latest_filing(cik, form_type)
        if not filing:
            print(f"    [skip] No {form_type} found for CIK {cik}.")
            return None, None
        acc_nodash, primary_doc = filing

    file_url = f"{EDGAR_ARCHIVE}/{cik}/{acc_nodash}/{primary_doc}"
    raw_ext = ("." + primary_doc.rsplit(".", 1)[-1].lower()) if "." in primary_doc else ".htm"

    # ── Non-HTML primary doc → needs manual placement ─────────────────────────
    if raw_ext in _MANUAL_EXTS:
        print(f"    [manual] Primary document is {raw_ext} — added to needed_files.md")
        return None, {
            "form_type": form_type,
            "company": company,
            "expected_filename": f"{filename}.pdf",
            "expected_path": str((save_dir / f"{filename}.pdf").relative_to(REPO_ROOT)),
            "edgar_url": file_url,
            "filing_index": f"https://www.sec.gov/cgi-bin/browse-edgar"
                            f"?action=getcompany&CIK={cik}&type={requests.utils.quote(form_type)}"
                            f"&dateb=&owner=include&count=5",
            "notes": f"Primary EDGAR document is {raw_ext.lstrip('.')} — please download a PDF "
                     f"version from the filing index and place it at the expected path above.",
        }

    # ── Download HTML and convert to PDF ─────────────────────────────────────
    raw_path = save_dir / f"{filename}{raw_ext}"
    try:
        resp = _edgar_get(file_url, stream=True, timeout=60)
        content_length = int(resp.headers.get("Content-Length", 0))
        if content_length and content_length > MAX_FILE_BYTES:
            print(f"    [skip] File too large ({content_length / 1e6:.1f} MB).")
            return None, None

        data = b""
        for chunk in resp.iter_content(chunk_size=65536):
            data += chunk
            if len(data) > MAX_FILE_BYTES:
                print(f"    [skip] Exceeded 20 MB during download.")
                return None, None

        raw_path.write_bytes(data)
        print(f"    [downloaded] {raw_path.name} ({len(data) / 1e3:.0f} KB) — converting to PDF …")

        if raw_ext == ".pdf":
            raw_path.rename(pdf_path)
            final_path = pdf_path
        else:
            final_path = _to_pdf(raw_path)
            if final_path:
                raw_path.unlink(missing_ok=True)
            else:
                final_path = raw_path   # keep raw on conversion failure

        size = final_path.stat().st_size
        print(f"    [ok] Saved {final_path.name} ({size / 1e3:.0f} KB)")
        return {
            "form_type": form_type, "company": company,
            "source_url": file_url,
            "file_path": str(final_path.relative_to(REPO_ROOT)),
            "format": final_path.suffix.lstrip("."),
            "size_bytes": size,
            "downloaded_at": datetime.now(timezone.utc).isoformat(),
        }, None

    except Exception as exc:
        print(f"    [error] Download failed: {exc}")
        return None, None


# ── FDIC BankFind downloader ───────────────────────────────────────────────────

def download_fdic_financials(source_url: str, save_dir: Path) -> Optional[Dict[str, Any]]:
    dest_path = save_dir / "fdic_financials.xlsx"
    if dest_path.exists():
        size = dest_path.stat().st_size
        print(f"  [exists] fdic_financials.xlsx — skipping")
        return {
            "form_type": "FFIEC Call Report (simulated)", "company": "FDIC BankFind",
            "source_url": "cached",
            "file_path": str(dest_path.relative_to(REPO_ROOT)),
            "format": "xlsx", "size_bytes": size,
            "downloaded_at": datetime.fromtimestamp(
                dest_path.stat().st_mtime, tz=timezone.utc
            ).isoformat(),
        }

    print("  Fetching FDIC financial data …")
    try:
        import pandas as pd
    except ImportError:
        print("    [skip] pandas not installed.")
        return None

    try:
        params = {
            "fields": "REPDTE,CERT,ASSET,DEP,INTINC,NONII,NETINC,EQ",
            "limit": 500, "offset": 0, "output": "json",
        }
        resp = _session.get(source_url, params=params, timeout=30)
        resp.raise_for_status()
        rows = [
            {k: v for k, v in item["data"].items() if k != "ID"}
            for item in resp.json().get("data", [])
        ]
        if not rows:
            print("    [warn] FDIC API returned no rows.")
            return None
        df = pd.DataFrame(rows)
        df.to_excel(dest_path, index=False, engine="openpyxl")
        size = dest_path.stat().st_size
        print(f"    [ok] Saved fdic_financials.xlsx ({size / 1e3:.0f} KB, {len(df)} rows)")
        return {
            "form_type": "FFIEC Call Report (simulated)", "company": "FDIC BankFind",
            "source_url": source_url,
            "file_path": str(dest_path.relative_to(REPO_ROOT)),
            "format": "xlsx", "size_bytes": size,
            "downloaded_at": datetime.now(timezone.utc).isoformat(),
        }
    except Exception as exc:
        print(f"    [error] FDIC download failed: {exc}")
        return None


# ── DOCX generator ─────────────────────────────────────────────────────────────

def generate_sample_docx(save_dir: Path) -> Optional[Dict[str, Any]]:
    dest_path = save_dir / "sample_kyc_document.docx"
    if dest_path.exists():
        size = dest_path.stat().st_size
        print(f"  [exists] sample_kyc_document.docx — skipping")
        return {
            "form_type": "KYC Onboarding Document (synthetic)", "company": "N/A",
            "source_url": "generated",
            "file_path": str(dest_path.relative_to(REPO_ROOT)),
            "format": "docx", "size_bytes": size,
            "downloaded_at": datetime.fromtimestamp(
                dest_path.stat().st_mtime, tz=timezone.utc
            ).isoformat(),
        }

    print("  Generating sample KYC DOCX …")
    try:
        from docx import Document
    except ImportError:
        print("    [skip] python-docx not installed.")
        return None

    doc = Document()
    doc.add_heading("KYC Onboarding Document", 0)
    doc.add_paragraph(
        "This document contains Know Your Customer (KYC) information "
        "for compliance verification. All data is synthetic."
    )
    doc.add_heading("1. Personal Information", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Field"
    tbl.rows[0].cells[1].text = "Value"
    for field, value in [
        ("Full Name", "Jane Elizabeth Doe"),
        ("Date of Birth", "1985-03-14"),
        ("Nationality", "United States of America"),
        ("Passport Number", "B04728193"),
        ("Passport Expiry", "2030-03-13"),
        ("Place of Birth", "New York, NY, USA"),
        ("Tax ID (SSN)", "XXX-XX-1234"),
        ("FBAR Filer", "Yes"),
    ]:
        r = tbl.add_row().cells
        r[0].text = field; r[1].text = value

    doc.add_heading("2. Address", level=1)
    doc.add_paragraph("Residential: 742 Evergreen Terrace, Springfield, IL 62704, USA")
    doc.add_heading("3. Beneficial Ownership", level=1)
    doc.add_paragraph(
        "UBO of Doe Holdings LLC (EIN: 84-1234567, Reg: LLC-2019-7734521), Delaware LLC."
    )
    doc.add_heading("4. Banking", level=1)
    doc.add_paragraph("Bank: First National Bank of Springfield")
    doc.add_paragraph("IBAN: DE89 3704 0044 0532 0130 00")
    doc.add_paragraph("IBAN: GB29 NWBK 6016 1331 9268 19  BIC: DEUTDEDB")
    doc.add_heading("5. Source of Funds", level=1)
    doc.add_paragraph("Employment income from Acme Corp, San Francisco CA. "
                       "Dividends from Acme Partners LP.")
    doc.add_heading("6. PEP / Sanctions", level=1)
    doc.add_paragraph("Subject is NOT a PEP and does NOT appear on OFAC/EU/UN lists.")

    doc.save(dest_path)
    size = dest_path.stat().st_size
    print(f"    [ok] Saved sample_kyc_document.docx ({size / 1e3:.0f} KB)")
    return {
        "form_type": "KYC Onboarding Document (synthetic)", "company": "N/A",
        "source_url": "generated",
        "file_path": str(dest_path.relative_to(REPO_ROOT)),
        "format": "docx", "size_bytes": size,
        "downloaded_at": datetime.now(timezone.utc).isoformat(),
    }


# ── Image fixtures ─────────────────────────────────────────────────────────────

def prepare_image_fixtures(save_dir: Path) -> List[Dict[str, Any]]:
    entries: List[Dict[str, Any]] = []
    src_png = TEST_DATA_ROOT / "passport.png"
    if not src_png.exists():
        print("    [warn] test-data/passport.png not found — skipping image fixtures.")
        return entries

    dest_png = save_dir / "passport.png"
    if dest_png.exists():
        print(f"    [exists] passport.png — skipping copy")
    else:
        shutil.copy2(src_png, dest_png)
        print(f"    [ok] Copied passport.png ({dest_png.stat().st_size / 1e3:.0f} KB)")
    entries.append({
        "form_type": "Passport scan (PNG)", "company": "N/A",
        "source_url": "test fixture",
        "file_path": str(dest_png.relative_to(REPO_ROOT)),
        "format": "png", "size_bytes": dest_png.stat().st_size,
        "downloaded_at": datetime.fromtimestamp(
            dest_png.stat().st_mtime, tz=timezone.utc
        ).isoformat(),
    })

    dest_tiff = save_dir / "passport.tiff"
    if dest_tiff.exists():
        print(f"    [exists] passport.tiff — skipping conversion")
    else:
        try:
            from PIL import Image
            Image.open(src_png).save(dest_tiff, format="TIFF")
            print(f"    [ok] Converted passport.tiff ({dest_tiff.stat().st_size / 1e3:.0f} KB)")
        except ImportError:
            print("    [skip] Pillow not installed — TIFF conversion skipped.")
            return entries
        except Exception as exc:
            print(f"    [warn] TIFF conversion failed: {exc}")
            return entries
    entries.append({
        "form_type": "Passport scan (TIFF)", "company": "N/A",
        "source_url": "test fixture",
        "file_path": str(dest_tiff.relative_to(REPO_ROOT)),
        "format": "tiff", "size_bytes": dest_tiff.stat().st_size,
        "downloaded_at": datetime.fromtimestamp(
            dest_tiff.stat().st_mtime, tz=timezone.utc
        ).isoformat(),
    })
    return entries


# ── Needed-files report ────────────────────────────────────────────────────────

def write_needed_files(needed: List[Dict[str, Any]]) -> None:
    if not needed:
        if NEEDED_PATH.exists():
            NEEDED_PATH.unlink()
        return

    lines = [
        "# Manually Required Test Files\n\n",
        "The following files could not be auto-downloaded because their primary "
        "EDGAR document is in XML or another non-HTML format.\n\n",
        "Please download a **PDF version** of each filing from the provided link "
        "and place it at the **expected path** relative to the repo root.\n\n",
    ]
    for i, e in enumerate(needed, 1):
        lines += [
            f"## {i}. {e['form_type']} — {e['company']}\n\n",
            f"- **Expected file:** `{e['expected_path']}`\n",
            f"- **EDGAR filing URL:** {e['edgar_url']}\n",
            f"- **Filing index:** {e['filing_index']}\n",
            f"- **Notes:** {e['notes']}\n\n",
        ]

    NEEDED_PATH.write_text("".join(lines))
    print(f"\nNeeded files report → {NEEDED_PATH.relative_to(REPO_ROOT)} ({len(needed)} entries)")


# ── Ingestion test runner ──────────────────────────────────────────────────────

_INGESTABLE_EXTS = {".pdf", ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".docx", ".xlsx"}


def _ingest_file(file_path: Path, api_url: str, customer_id: str, form_type: str) -> Dict[str, Any]:
    ingest_url = f"{api_url.rstrip('/')}/api/v1/ingest"
    result: Dict[str, Any] = {
        "form_type": form_type, "file": file_path.name,
        "status": "error", "doc_id": None,
        "chunks": None, "elapsed_ms": None, "error": None,
    }
    t0 = time.perf_counter()
    try:
        with open(file_path, "rb") as fh:
            resp = requests.post(
                ingest_url,
                files={"file": (file_path.name, fh)},
                data={"customer_id": customer_id},
                timeout=30,
            )
        if resp.status_code not in (200, 202):
            result["error"] = f"HTTP {resp.status_code}: {resp.text[:200]}"
            result["elapsed_ms"] = round((time.perf_counter() - t0) * 1000)
            return result
        result["doc_id"] = resp.json().get("doc_id")
    except Exception as exc:
        result["error"] = str(exc)
        result["elapsed_ms"] = round((time.perf_counter() - t0) * 1000)
        return result

    status_url = f"{api_url.rstrip('/')}/api/v1/ingest/{result['doc_id']}/status"
    deadline = time.time() + 300
    while time.time() < deadline:
        time.sleep(3)
        try:
            sdata = requests.get(status_url, timeout=15).json()
            st = sdata.get("status", "queued")
            if st == "completed":
                result["status"] = "completed"
                result["chunks"] = sdata.get("semantic_chunks_created")
                break
            elif st == "failed":
                result["status"] = "failed"
                result["error"] = sdata.get("error")
                break
        except Exception:
            pass
    else:
        result["status"] = "timeout"
        result["error"] = "Pipeline did not complete within 5 minutes."

    result["elapsed_ms"] = round((time.perf_counter() - t0) * 1000)
    return result


def run_ingestion_tests(manifest: List[Dict[str, Any]], api_url: str, customer_id: str) -> None:
    print(f"\n{'─' * 80}")
    print(f"Ingestion smoke tests  →  {api_url}  (customer: {customer_id})")
    print(f"{'─' * 80}\n")
    results = []
    for entry in manifest:
        fp = REPO_ROOT / entry["file_path"]
        if not fp.exists():
            print(f"  [skip] {entry['file_path']} not found.")
            continue
        if fp.suffix.lower() not in _INGESTABLE_EXTS:
            print(f"  [skip] {fp.name} — unsupported format '{entry['format']}'.")
            continue
        print(f"  Ingesting {fp.name} ({entry['form_type']}) …")
        r = _ingest_file(fp, api_url, customer_id, entry["form_type"])
        results.append(r)
        print(f"    → {r['status'].upper()} | doc_id={r['doc_id']} | "
              f"chunks={r['chunks']} | {r['elapsed_ms']} ms")
        if r["error"]:
            print(f"       error: {r['error']}")

    print(f"\n{'─' * 80}")
    print(f"{'Form Type':<35} {'File':<30} {'Status':<12} {'Chunks':<8} ms")
    print(f"{'─' * 80}")
    for r in results:
        print(f"{r['form_type']:<35} {r['file']:<30} {r['status']:<12} "
              f"{str(r['chunks'] or ''):<8} {r['elapsed_ms'] or ''}")
    print(f"{'─' * 80}")
    passed = sum(1 for r in results if r["status"] == "completed")
    print(f"\nResults: {passed} passed, {len(results) - passed} failed, {len(results)} total\n")


# ── Directory setup ────────────────────────────────────────────────────────────

def _setup_dirs() -> None:
    for d in (SEC_FILINGS_DIR, BANKING_DIR, IMAGES_DIR, DOCUMENTS_DIR):
        d.mkdir(parents=True, exist_ok=True)


# ── Main ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Download test documents (driven by document_catalog.xlsx) "
                    "and optionally run ingestion smoke tests."
    )
    parser.add_argument("--api-url", default="http://localhost:8000")
    parser.add_argument("--customer-id", default="test-co-001")
    parser.add_argument("--skip-download", action="store_true",
                        help="Skip downloads; use already-downloaded files.")
    parser.add_argument("--skip-ingest", action="store_true",
                        help="Only download; do not call the ingest API.")
    parser.add_argument("--forms", default="",
                        help="Comma-separated form types to process (default: all).")
    args = parser.parse_args()

    form_filter = {f.strip() for f in args.forms.split(",") if f.strip()}
    _setup_dirs()
    manifest: List[Dict[str, Any]] = []
    needed:   List[Dict[str, Any]] = []

    # ── Phase 1: Downloads ─────────────────────────────────────────────────────

    if not args.skip_download:
        catalog = load_catalog()
        print(f"\nLoaded {len(catalog)} entries from {CATALOG_PATH.name}\n")

        _images_done = False   # generate image fixtures only once

        for row in catalog:
            ft = row["form_type"]
            if form_filter and ft not in form_filter:
                continue

            source_url = row.get("source_url", "")
            is_manual  = row["manual_download"]

            # ── FDIC row ────────────────────────────────────────────────────
            if "banks.data.fdic.gov" in source_url:
                print(f"=== FDIC row: {ft} ===")
                r = download_fdic_financials(source_url, BANKING_DIR)
                if r:
                    manifest.append(r)
                continue

            # ── Generated DOCX row ──────────────────────────────────────────
            if source_url == "generated":
                print(f"=== Synthetic DOCX: {ft} ===")
                r = generate_sample_docx(DOCUMENTS_DIR)
                if r:
                    manifest.append(r)
                continue

            # ── Image fixture rows ───────────────────────────────────────────
            if source_url == "test fixture":
                if not _images_done:
                    print(f"=== Image fixtures ===")
                    manifest.extend(prepare_image_fixtures(IMAGES_DIR))
                    _images_done = True
                continue

            # ── Manual-download row ──────────────────────────────────────────
            cik = row.get("cik", "").strip()
            company = row.get("company", "")
            filename = row.get("filename", ft.replace(" ", "_"))

            if is_manual:
                # Resolve the EDGAR URL for the user's reference, but don't download
                print(f"  [manual] {ft} ({company}) — looking up filing URL …")
                acc_nodash = primary_doc = None
                try:
                    if cik == "0":
                        result = _efts_find(ft)
                        if result:
                            cik_r, acc_nodash, company = result
                            primary_doc = _get_primary_doc_from_submissions(cik_r, acc_nodash)
                            file_url = (f"{EDGAR_ARCHIVE}/{cik_r}/{acc_nodash}/{primary_doc}"
                                        if primary_doc else "")
                            index_url = (
                                f"https://www.sec.gov/cgi-bin/browse-edgar"
                                f"?action=getcompany&CIK={cik_r}"
                                f"&type={requests.utils.quote(ft)}&dateb=&owner=include&count=5"
                            )
                        else:
                            file_url = f"https://www.sec.gov/cgi-bin/browse-edgar?action=getcompany&type={requests.utils.quote(ft)}&dateb=&owner=include&count=5"
                            index_url = file_url
                    else:
                        filing = _get_latest_filing(cik, ft)
                        if filing:
                            acc_nodash, primary_doc = filing
                            file_url = f"{EDGAR_ARCHIVE}/{cik}/{acc_nodash}/{primary_doc}"
                            index_url = (
                                f"https://www.sec.gov/cgi-bin/browse-edgar"
                                f"?action=getcompany&CIK={cik}"
                                f"&type={requests.utils.quote(ft)}&dateb=&owner=include&count=5"
                            )
                        else:
                            file_url = ""
                            index_url = f"https://efts.sec.gov/LATEST/search-index?forms={requests.utils.quote(ft)}"
                except Exception:
                    file_url = ""
                    index_url = f"https://efts.sec.gov/LATEST/search-index?forms={requests.utils.quote(ft)}"

                raw_ext = ("." + primary_doc.rsplit(".", 1)[-1].lower()) if primary_doc and "." in primary_doc else ".xml"
                needed.append({
                    "form_type": ft,
                    "company": company,
                    "expected_filename": f"{filename}.pdf",
                    "expected_path": str((SEC_FILINGS_DIR / f"{filename}.pdf").relative_to(REPO_ROOT)),
                    "edgar_url": file_url,
                    "filing_index": index_url,
                    "notes": (
                        f"Primary EDGAR document is {raw_ext.lstrip('.')} — "
                        "download a PDF version from the filing index and place it at the expected path."
                    ),
                })
                print(f"    [manual] Added to needed_files.md")
                continue

            # ── Auto EDGAR download ─────────────────────────────────────────
            manifest_entry, needed_entry = download_from_edgar(
                form_type=ft,
                cik=cik,
                company=company,
                filename=filename,
                save_dir=SEC_FILINGS_DIR,
            )
            if manifest_entry:
                manifest.append(manifest_entry)
            if needed_entry:
                needed.append(needed_entry)

        MANIFEST_PATH.write_text(json.dumps(manifest, indent=2))
        print(f"\nManifest written → {MANIFEST_PATH.relative_to(REPO_ROOT)} ({len(manifest)} entries)")
        write_needed_files(needed)

    else:
        if MANIFEST_PATH.exists():
            manifest = json.loads(MANIFEST_PATH.read_text())
            print(f"\nLoaded manifest ({len(manifest)} entries)")
        else:
            print("[warn] --skip-download set but manifest.json not found.")
            sys.exit(0)

    # ── Phase 2: Ingestion tests ───────────────────────────────────────────────

    if not args.skip_ingest and manifest:
        run_ingestion_tests(manifest, args.api_url, args.customer_id)
    elif args.skip_ingest:
        print("\n[info] --skip-ingest: ingestion tests skipped.")


if __name__ == "__main__":
    main()
